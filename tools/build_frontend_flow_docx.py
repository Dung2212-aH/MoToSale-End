from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "frontend-backend-database-flows.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_note(doc, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_cell_margins(table, top=100, bottom=100, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(body)


def add_code_block(doc, title: str, lines: list[str]) -> None:
    add_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_cell_margins(table, top=100, bottom=100, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor(45, 45, 45)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_cell_margins(table)
    hdr = table.rows[0]
    hdr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(11, 37, 69)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        s = styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        s = styles[list_style]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25


PAGE_ROUTES = [
    ("/", "HomePage.jsx + Header", "Load san pham noi bat, danh muc, bo loc header, don hang gan day neu da dang nhap.", "productApi.getAll; categoryApi.getAll; productApi.getFilters; orderApi.getMyOrders; voucherApi.getMineCount", "GET /products; GET /categories; GET /products/filters; GET /orders; GET /vouchers/my/count", "CatalogService, OrderService"),
    ("/products", "ProductListPage.jsx", "Danh sach san pham, loc theo danh muc/hang/gia, sap xep, phan trang, them gio hang nhanh.", "productApi.getFilters; productApi.getAll; productApi.getById khi quick-add", "GET /products/filters; GET /products; GET /products/{id}", "CatalogService"),
    ("/products/:id", "ProductDetailPage.jsx + ProductInfoBox + ProductReviews", "Chi tiet san pham, bien the, anh, san pham cung hang/cung loai, yeu thich, danh gia, them gio hang.", "productApi.getById; productApi.getAll; reviewApi.*; favoriteApi.*; cartApi.addItem", "GET /products/{id}; GET /products; GET/POST /products/{id}/reviews; GET /reviews/product/{id}/me; GET/POST/DELETE /favorites; POST /cart/items", "CatalogService, OrderService"),
    ("/cart", "CartPage.jsx + CartContext", "Xem gio hang, tang/giam so luong, xoa dong, xoa toan bo.", "cartApi.getMine; addItem; updateItem; removeItem; clear", "GET /cart; POST /cart/items; PUT /cart/items/{id}; DELETE /cart/items/{id}; DELETE /cart/clear", "OrderService"),
    ("/checkout", "CheckoutPage.jsx", "Lay dia chi, tinh voucher/phi van chuyen, tao don, giu ton kho, tao thong tin thanh toan/dat coc/tra gop.", "userApi.getAddresses; voucherApi.getApplicable/validate; orderApi.getShippingQuote/create", "GET /users/me/addresses; POST /vouchers/applicable; POST /vouchers/validate; POST /orders/shipping-quote; POST /orders", "AuthService, OrderService"),
    ("/checkout/payment", "PaymentPage.jsx", "Hien thi thong tin thanh toan don, huy don truoc thanh toan.", "orderApi.getById; orderApi.getPaymentInfo; orderApi.cancel", "GET /orders/{id}; GET /orders/{id}/payment-info; PUT /orders/{id}/cancel", "OrderService"),
    ("/checkout/success", "CheckoutSuccessPage.jsx", "Doc lai don sau khi tao/thanh toan.", "orderApi.getById", "GET /orders/{id}", "OrderService"),
    ("/orders", "OrdersPage.jsx", "Danh sach don cua khach, dem don cho thanh toan, kiem tra quyen danh gia san pham da mua.", "orderApi.getMyOrders; orderApi.getAll; reviewApi.getMine", "GET /orders; GET /reviews/product/{productId}/me", "OrderService, CatalogService"),
    ("/orders/:id", "OrderDetailPage.jsx", "Chi tiet don, lich su, thanh toan, huy don, yeu cau hoan tien.", "orderApi.getById; paymentApi.getByOrder; orderApi.getPaymentInfo; orderApi.cancel; orderApi.requestRefund", "GET /orders/{id}; GET /payments/order/{orderId}; GET /orders/{id}/payment-info; PUT /orders/{id}/cancel; POST /orders/{id}/request-refund", "OrderService, PaymentService"),
    ("/account", "AccountPage.jsx", "Ho so ca nhan, doi mat khau, CRUD dia chi giao hang.", "userApi.getProfile/updateProfile/changePassword; address APIs", "GET/PUT /users/me; PUT /users/me/password; GET/POST/PUT/DELETE /users/me/addresses", "AuthService"),
    ("/favorites", "FavoritesPage.jsx + FavoriteContext", "Danh sach yeu thich, them/bo yeu thich, quick-add san pham.", "favoriteApi.getMine/add/remove; productApi.getById; cartApi.addItem", "GET /favorites; POST /favorites/{productId}; DELETE /favorites/{productId}; GET /products/{id}; POST /cart/items", "CatalogService, OrderService"),
    ("/vouchers", "VouchersPage.jsx", "Xem voucher dang hieu luc, voucher da luu, luu voucher.", "voucherApi.getAll; getMine; save", "GET /vouchers/active; GET /vouchers/my; POST /vouchers/save", "OrderService"),
    ("/contact", "ContactPage.jsx", "Gui yeu cau lien he/tu van.", "contentApi.createContactRequest", "POST /content/contact-requests", "CatalogService"),
    ("/faq", "FaqPage.jsx", "Hien thi cau hoi thuong gap.", "contentApi.getFaqs", "GET /content/faqs", "CatalogService"),
    ("/login", "LoginPage.jsx + AuthContext", "Dang nhap, luu JWT/session, gan Authorization bearer cho cac request sau.", "authApi.login", "POST /auth/login", "AuthService"),
    ("/register", "RegisterPage.jsx + AuthContext", "Dang ky tai khoan khach hang.", "authApi.register", "POST /auth/register", "AuthService"),
    ("/forgot-password", "ForgotPasswordPage.jsx", "Tao token dat lai mat khau va dat mat khau moi.", "authApi.forgotPassword; authApi.resetPassword", "POST /auth/forgot-password; POST /auth/reset-password", "AuthService"),
]


FRONTEND_CODE_MAP = [
    ("Frontend/src/App.jsx", "Composition root cua storefront.", "Khai bao BrowserRouter, providers va route tree. AuthProvider boc ngoai Cart/Favorite de cac context con biet user da dang nhap chua. PublicRoute chan user da login vao login/register; ProtectedRoute chan route can token."),
    ("Frontend/src/components/ProtectedRoute.jsx", "Gate bao ve route.", "Doc isAuthenticated/loading tu AuthContext. Neu dang loading thi render spinner; neu chua co token thi redirect sang /login?redirect=<current-url>; neu hop le thi render children."),
    ("Frontend/src/services/httpClient.js", "Lop HTTP thap nhat.", "Tao axios instance voi baseURL /api. Quan ly token trong sessionStorage/localStorage, decode JWT de lay claim, request interceptor gan Authorization: Bearer <token>, response interceptor day message backend vao error.message."),
    ("Frontend/src/services/api.js", "Facade API nghiep vu.", "Gom cac nhom authApi/productApi/cartApi/orderApi/... Moi ham goi endpoint backend, map payload UI sang DTO backend tieng Viet, sau do normalize response ve shape frontend can."),
    ("Frontend/src/services/normalizers.js", "Chuan hoa response va query.", "field() lay gia tri theo nhieu alias camel/Pascal/tieng Viet; toQuery() doi filter UI thanh query backend; mapOrder/mapVoucher/mapPayment/mapFavorite/mapReview gom nhieu response shape ve 1 object on dinh."),
    ("Frontend/src/utils/productMappers.js", "Mapper san pham/gio hang.", "normalizeProduct doc maSanPham/tenSanPham/giaGoc/giaKhuyenMai/giaBan/tyLeGiam/bienThe/anh va doi thanh id/name/basePrice/salePrice/discountPercent/variants/images. normalizeCart map cart item, gia, anh, variant."),
    ("Frontend/src/utils/productOptions.js", "Tao option bien the cho detail.", "Tu variants/images suy ra versions, colors, selected image mapping, color fallback. Day la lop UI logic, khong goi API."),
    ("Frontend/src/contexts/AuthContext.jsx", "State phien dang nhap.", "Khoi tao user tu authApi.getCurrentUser(); login/register/logout; expose isAuthenticated. Lang nghe storage va AUTH_CHANGED_EVENT de dong bo tab/window."),
    ("Frontend/src/contexts/CartContext.jsx", "State gio hang dung chung.", "Khi authenticated thi refreshCart GET /cart. add/update/remove/clear goi cartApi roi applyCart cap nhat count va mini cart. Kiem tra stock local truoc khi update quantity."),
    ("Frontend/src/contexts/FavoriteContext.jsx", "State yeu thich.", "loadFavorites GET /favorites khi login; isFavorite dung Set id; toggleFavorite optimistic update UI roi POST/DELETE backend, rollback neu loi."),
    ("Frontend/vite.config.js", "Dev proxy.", "Server frontend port 5174. Proxy /api va /uploads sang http://localhost:5000, tuc API Gateway; frontend code khong goi truc tiep 5001/5002/5003/5004."),
]


API_CONTRACTS = [
    ("authApi.login", "Input UI: username, password, rememberMe.", "POST /auth/login body { email: username, matKhau: password }.", "normalizeLoginResponse lay token/user/role; saveAuthUser luu session/local theo rememberMe."),
    ("authApi.register", "Input UI: name, email, phone, password.", "POST /auth/register body { hoTen, email, soDienThoai, matKhau }.", "Backend tao khach hang; frontend thuong dieu huong ve login."),
    ("productApi.getAll", "Input UI: keyword, categoryId, brandId, minPrice, maxPrice, sortBy, page.", "GET /products params gom DangHoatDong=true va toQuery map categoryId->MaDanhMuc, brandId->MaHangXe, minPrice->GiaTu, sortBy->SortBy/SortDescending.", "normalizeProductList tra items/page/pageSize/totalPages; moi item thanh ProductCard model."),
    ("productApi.getById", "Input: id tren URL hoac id san pham quick-add.", "GET /products/{id}.", "normalizeProduct gom thong tin san pham, variants, images; ProductDetailPage tiep tuc normalizeProductOptions de tao version/color."),
    ("reviewApi.create", "Input UI: rating, comment, title, image, optional orderId.", "POST multipart /products/{productId}/reviews. buildReviewForm map rating->Diem, comment->NoiDung, image->Image.", "Backend luu review Pending; summary storefront chi tinh Approved."),
    ("cartApi.addItem", "Input UI: productId, variantId, quantity.", "POST /cart/items body { maSanPham, maBienSanPham, soLuong }.", "handleCart normalizeCart va notifyCartChanged de Header/CartContext cap nhat count."),
    ("cartApi.updateItem", "Input UI: cart item id, quantity.", "PUT /cart/items/{id} body { soLuong }.", "Sau PUT goi lai GET /cart de co snapshot moi tu server."),
    ("orderApi.create", "Input CheckoutPage payload: shipping, receivingMethod, orderType, paymentMethod, voucherCode, installment/deposit.", "POST /orders body dung ten field backend: hoTenNhanHang, soDienThoaiNhanHang, diaChiNhanHang, maVoucherCode, phuongThucNhanHang, loaiDonHang, phuongThucThanhToan, hoSoTraGop, tienDatCoc.", "mapOrder tra order id/code/status; frontend navigate payment/success tuy payment method."),
    ("orderApi.getShippingQuote", "Input: receivingMethod, shippingProvince, voucherCode, orderType.", "POST /orders/shipping-quote body { phuongThucNhanHang, shippingProvince, maVoucherCode, orderType }.", "Tra shippingFee/originalShippingFee/discountAmount/carrierName cho checkout preview."),
    ("voucherApi.validate", "Input UI co subtotal/productIds/categoryIds/brandIds/orderType/shippingFee.", "POST /vouchers/validate. Backend thuc te recompute tu active cart va stored procedure; frontend gui them context de tuong thich UI.", "Tra valid/message/discountAmount/voucher; loi 400 duoc interceptor day message len UI."),
    ("userApi.createAddress/updateAddressById", "Input UI: fullName, phoneNumber, addressLine, ward, district, province, note, isDefault.", "mapAddressBody -> hoTenNhanHang, soDienThoaiNhanHang, diaChiNhanHang, ward, district, province, ghiChu, laMacDinh.", "Backend luu NGUOIDUNG_DIACHI; checkout doc lai address list."),
    ("favoriteApi.add/remove", "Input: product id.", "POST /favorites/{productId} hoac DELETE /favorites/{productId}.", "FavoriteContext optimistic update UI, backend dam bao unique theo user + product."),
    ("contentApi.createContactRequest", "Input contact form: fullName, phoneNumber, email, subject, message, inquiryType, productId.", "POST /content/contact-requests body { hoTen, soDienThoai, email, tieuDe, noiDung, loaiYeuCau, maSanPham }.", "Backend insert LIENHE_YEUCAU va co the gan MaSanPham neu ton tai."),
]


PAGE_CODE_DETAILS = [
    ("HomePage.jsx", "useEffect load productApi.getAll + categoryApi.getAll; neu login thi load orderApi.getMyOrders. Header rieng load productApi.getFilters va voucher count.", "Trang home khong giu logic gia/DB; no nhan san pham da normalize roi chia thanh section deal/bestseller/category. Quick add san pham co bien the se getById de buoc user vao detail chon SKU."),
    ("ProductListPage.jsx", "State filter/sort/page lay tu URL/search params. build apiQueryValues roi Promise.all(productApi.getFilters, productApi.getAll).", "toQuery trong normalizers.js la diem map quan trong: UI dung categoryId/brandId/minPrice, backend nhan MaDanhMuc/MaHangXe/GiaTu. ProductCard chi render model normalized."),
    ("ProductDetailPage.jsx", "useAsync(productApi.getById(id)); normalizeProductOptions(product) tao variants/colors/images; selectedVariant duoc tinh bang matchesSelection(version,color).", "Trang nay ghep catalog + cart + favorite + review: addToCart tao payload productId/variantId/quantity; related sections goi productApi.getAll theo brand/category; viewed products luu local storage."),
    ("ProductReviews.jsx / ReviewModal.jsx", "Load song song reviewApi.getByProduct, getSummary, getMine. Submit goi reviewApi.create multipart.", "Backend chi hien Approved trong danh sach. getMine tra canReview/hasPurchased/eligibleOrderId de UI biet co cho form danh gia hay khong."),
    ("CartPage.jsx + CartContext.jsx", "CartContext refreshCart khi isAuthenticated thay doi. CartPage thao tac update/remove/clear qua context.", "Context la source of truth cho cart count. Moi mutation backend tra cart moi hoac frontend goi lai GET /cart de tranh lech tong tien/stock."),
    ("CheckoutPage.jsx", "Load user addresses, tinh subtotal tu cart context, goi voucher applicable/validate va shipping quote, build payload create order.", "Day la luong nhieu contract nhat: UI field tieng Anh duoc map trong orderApi.create sang field tieng Viet. Backend moi la noi tinh gia cuoi, voucher, shipping, hold stock."),
    ("PaymentPage.jsx", "Doc orderId tu URL/state, load order detail + payment-info, hien QR/chuyen khoan, cho huy don neu chua thanh toan.", "Huy don goi OrderService, khong xoa don; backend cap nhat status, lich su, ton giu cho va voucher."),
    ("CheckoutSuccessPage.jsx", "Doc orderId va goi orderApi.getById.", "Trang xac nhan chi doc data; khong tao side effect database moi."),
    ("OrdersPage.jsx", "Load orderApi.getMyOrders va mot query AwaitingPayment. Co kiem tra reviewApi.getMine theo product khi user muon danh gia.", "mapOrder gom order/detail/voucher/tra gop ve shape UI. Review eligibility doc CatalogService vi review nam o catalog context."),
    ("OrderDetailPage.jsx", "Load order detail, payments by order, payment-info. Handlers cancel va requestRefund.", "requestRefund tao YEUCAU_HOANTIEN; paymentApi.getByOrder doc PaymentService, con payment-info doc OrderService de lay cau hinh ngan hang."),
    ("AccountPage.jsx", "Promise.all userApi.getProfile + getAddresses. Handlers update profile, change password, CRUD address.", "Tat ca di AuthService/UsersController; address co fallback endpoint cu /users/me/address neu /addresses 404."),
    ("FavoritesPage.jsx + FavoriteContext.jsx", "Context load favorites khi login; page render favoriteProducts. Quick-add cung pattern ProductList: getById neu can variant.", "Favorite item backend co product nested; normalizeProduct dam bao ProductCard dung chung model."),
    ("VouchersPage.jsx", "Load voucherApi.getAll active; neu login load voucherApi.getMine. Save goi voucherApi.save.", "Voucher hien thi public active tu OrderService; voucher cua user la bang VOUCHER_NGUOIDUNG."),
    ("ContactPage.jsx", "Validate form client-side, submit contentApi.createContactRequest.", "Backend insert LIENHE_YEUCAU; neu co productId thi lien ket san pham de staff xu ly sau."),
    ("FaqPage.jsx", "Load contentApi.getFaqs va loc/search tren client.", "Backend tra FAQ active; frontend khong mutate."),
    ("Login/Register/ForgotPassword", "Auth pages goi AuthContext/authApi. Login thanh cong save token va dispatch AUTH_CHANGED_EVENT.", "AUTH_CHANGED_EVENT lam CartContext/FavoriteContext/Header dong bo theo user moi."),
]


BACKEND_CODE_MAP = [
    ("ApiGateway/ocelot.json", "Routing layer.", "Khong xu ly nghiep vu; no map upstream /api/... tu frontend sang service port noi bo. Khi debug loi 404/405 can xem route gateway truoc controller."),
    ("AuthService/Controllers/AuthController.cs", "Endpoint dang nhap/dang ky/reset.", "Controller mong request body tieng Viet tu api.js. No goi IAuthService, tra token/user hoac message loi."),
    ("AuthService/Controllers/UsersController.cs", "Endpoint /users/me va /users/me/addresses.", "Lay user id tu JWT claims, doc/ghi NGUOIDUNG va NGUOIDUNG_DIACHI. Cac endpoint admin cung trong file nhung khong thuoc storefront scope."),
    ("AuthService/Services/AuthService.cs", "Business logic auth.", "Validate password, hash password, tao JWT, tao/reset token. Repository che giau EF query truc tiep."),
    ("CatalogService/Controllers/ProductsController.cs", "Catalog endpoints.", "GET /products va /products/{id} goi CatalogService; /filters tu DbContext. Cac endpoint CRUD/variants/images la admin, khong tinh storefront tru khi detail doc BienThe/Anh."),
    ("CatalogService/Services/CatalogService.cs", "Mapper catalog.", "GetProductsAsync lay page product, imageMap, brand/category maps, review summary roi MapProductListItem. GetProductByIdAsync lay variants/images va MapProductDetail."),
    ("CatalogService/Repositories/Products/ProductRepository.cs", "Product query builder.", "Nhan ProductSearchDto, apply keyword/category/brand/model/type/status/price/sort/page. Day la noi filter ProductListPage tac dong xuong SQL."),
    ("CatalogService/Controllers/ReviewsController.cs", "Review endpoints.", "Public GET chi tra Approved. Authenticated getMine/create dung order history de kiem tra user da mua truoc khi review."),
    ("CatalogService/Controllers/FavoritesController.cs", "Favorite endpoints.", "User id lay tu JWT; GET join YEUTHICH voi SANPHAM; POST dam bao product ton tai va favorite chua co; DELETE xoa theo user/product."),
    ("CatalogService/Controllers/ContentController.cs", "FAQ/contact.", "GET /content/faqs doc FAQ active. POST /content/contact-requests validate input va insert LIENHE_YEUCAU."),
    ("OrderService/Controllers/CartController.cs", "Cart endpoints.", "Controller mong JWT user. Moi action goi OrderService roi tra CartDto; frontend khong tinh authoritative total."),
    ("OrderService/Services/OrderService.cs", "Core order/cart business logic.", "Xu ly cart, checkout, shipping quote, payment-info, cancel, refund. Day la lop quan trong nhat cho side effect database cua storefront."),
    ("OrderService/Repositories/OrderRepository.cs", "Order persistence + stored procedures.", "Lay order detail/search, tao order tu cart, apply inventory hold, validate/record/cancel voucher bang stored procedure."),
    ("OrderService/Controllers/VouchersController.cs", "Voucher endpoints.", "GET active/my/count/save/applicable/validate. applicable co raw SQL vao bang scope voucher; validate goi sp_Voucher_KiemTraTruocKhiTaoDon."),
    ("PaymentService/Controllers/PaymentsController.cs", "Payment read endpoint cho storefront.", "Frontend chi goi GET /payments/order/{orderId}. Cac endpoint confirm/fail/cancel chu yeu cho admin/payment workflow."),
    ("PaymentService/Repositories/PaymentRepository.cs", "Payment persistence.", "Doc THANHTOAN theo order; mot so action payment co the goi sp_TONKHO_ApDungBienDong khi confirm thanh toan/ton."),
]


DB_TABLE_DETAILS = [
    ("NGUOIDUNG", "AuthService, OrderService, PaymentService, CatalogService review", "Tai khoan, thong tin profile, customer id trong order/review/payment."),
    ("NGUOIDUNG_DIACHI", "AuthService, OrderService", "Dia chi giao/nhan hang cua user; checkout doc qua AuthService, order snapshot dia chi vao DONHANG."),
    ("SANPHAM", "CatalogService, OrderService, PaymentService", "Nguon catalog san pham; list/detail doc tai CatalogService, cart/order doc lai ten/gia/trang thai."),
    ("BIENSANPHAM", "CatalogService, OrderService, PaymentService", "SKU/bien the, ton kho, gia ghi de hien tai; add cart/order dung MaBienSanPham neu co."),
    ("ANHSANPHAM", "CatalogService, OrderService", "Anh chinh/anh bien the cho product card/detail/cart."),
    ("DANHMUC/HANGXE/DONGXE", "CatalogService", "Bo loc san pham va ten hien thi category/brand/model."),
    ("GIOHANG/CHITIET_GIOHANG", "OrderService", "Cart active cua user; checkout tao DONHANG tu cac dong nay."),
    ("DONHANG/CHITIET_DONHANG", "OrderService, CatalogService review", "Order header/detail. CatalogService doc de xac minh quyen danh gia; OrderService la owner nghiep vu don."),
    ("DONHANG_LICHSU_TRANGTHAI", "OrderService", "Audit trang thai don: tao, huy, thanh toan, fulfillment."),
    ("TONKHO_GIUCHO", "OrderService, PaymentService", "Giu ton cho don chua hoan tat; duoc giai phong khi huy/het han/tuy payment flow."),
    ("THANHTOAN", "OrderService, PaymentService", "Giao dich/yeu cau thanh toan; storefront doc qua payment info va payment list."),
    ("HOSO_TRAGOP", "OrderService", "Thong tin ho so tra gop khi checkout orderType/paymentMethod yeu cau."),
    ("YEUCAU_HOANTIEN", "OrderService", "Refund request do khach tao tu OrderDetailPage."),
    ("VOUCHER", "OrderService", "Dinh nghia ma giam gia, type/scope/date/usage."),
    ("VOUCHER_NGUOIDUNG", "OrderService", "Voucher user da luu va usage theo user."),
    ("DONHANG_VOUCHER", "OrderService", "Snapshot voucher da ap dung vao don."),
    ("VOUCHER_SANPHAM/VOUCHER_DANHMUC/VOUCHER_HANGXE", "OrderService raw SQL", "Bang scope ap dung voucher theo san pham/danh muc/hang."),
    ("DANHGIASANPHAM", "CatalogService", "Review cua user; public list chi lay Approved."),
    ("YEUTHICH", "CatalogService", "Favorite theo user + product."),
    ("FAQ", "CatalogService", "FAQ storefront."),
    ("LIENHE_YEUCAU", "CatalogService", "Yeu cau lien he/tu van tu ContactPage."),
    ("MATKHAU_DATLAI", "AuthService", "Token reset password."),
]


ENDPOINTS = [
    ("authApi.login", "POST /auth/login", "AuthController.Login", "AuthService + UserRepository", "NGUOIDUNG, VAITRO, NGUOIDUNG_VAITRO", "Doc user theo email, kiem tra mat khau, role; tra JWT."),
    ("authApi.register", "POST /auth/register", "AuthController.Register", "AuthService + UserRepository", "NGUOIDUNG, VAITRO, NGUOIDUNG_VAITRO", "Kiem tra email trung, tao user, gan role khach hang."),
    ("authApi.forgotPassword", "POST /auth/forgot-password", "AuthController.ForgotPassword", "AuthService", "NGUOIDUNG, MATKHAU_DATLAI", "Tao token reset password; gui/tra thong tin theo cau hinh."),
    ("authApi.resetPassword", "POST /auth/reset-password", "AuthController.ResetPassword", "AuthService", "NGUOIDUNG, MATKHAU_DATLAI", "Xac thuc token, cap nhat hash mat khau, danh dau token da dung."),
    ("productApi.getAll", "GET /products", "ProductsController.GetProducts", "CatalogService.GetProductsAsync + ProductRepository", "SANPHAM, ANHSANPHAM, HANGXE, DANHMUC, DANHGIASANPHAM", "Loc/sap xep/phan trang san pham dang hoat dong; lay anh chinh, ten hang/danh muc, tong hop danh gia."),
    ("productApi.getById", "GET /products/{id}", "ProductsController.GetProductById", "CatalogService.GetProductByIdAsync", "SANPHAM, BIENSANPHAM, ANHSANPHAM, HANGXE, DANHMUC, DANHGIASANPHAM", "Lay chi tiet san pham, bien the, anh, diem/tong danh gia da duyet."),
    ("productApi.getFilters", "GET /products/filters", "ProductsController.GetFilters", "CatalogDbContext LINQ", "DANHMUC, HANGXE, DONGXE", "Lay bo loc danh muc, hang xe, dong xe dang hoat dong."),
    ("categoryApi.getAll", "GET /categories", "CategoriesController.GetCategories", "CatalogService.GetCategoriesAsync", "DANHMUC", "Lay danh muc dang hoat dong."),
    ("reviewApi.getByProduct", "GET /products/{id}/reviews", "ReviewsController.GetByProduct", "Catalog DbContext", "DANHGIASANPHAM, NGUOIDUNG", "Lay danh gia Approved cua san pham."),
    ("reviewApi.getSummary", "GET /products/{id}/reviews/summary", "ReviewsController.GetSummary", "Catalog DbContext", "DANHGIASANPHAM", "Group by san pham de tinh so luong va diem trung binh."),
    ("reviewApi.getMine", "GET /reviews/product/{id}/me", "ReviewsController.GetMineForProduct", "Catalog DbContext", "DANHGIASANPHAM, DONHANG, CHITIET_DONHANG", "Kiem tra user da mua san pham trong don du dieu kien va danh gia hien co."),
    ("reviewApi.create", "POST /products/{id}/reviews", "ReviewsController.Create", "Catalog DbContext + image storage", "DANHGIASANPHAM, DONHANG, CHITIET_DONHANG", "Kiem tra quyen danh gia, luu danh gia Pending va anh neu co."),
    ("cartApi.getMine", "GET /cart", "CartController.GetMine", "OrderService.GetCartAsync", "GIOHANG, CHITIET_GIOHANG, SANPHAM, BIENSANPHAM, ANHSANPHAM", "Lay gio hang active cua user, dong san pham, bien the, gia va anh."),
    ("cartApi.addItem", "POST /cart/items", "CartController.AddItem", "OrderService.AddCartItemAsync", "GIOHANG, CHITIET_GIOHANG, SANPHAM, BIENSANPHAM", "Tao gio hang neu chua co; kiem tra ton/gia; upsert dong gio hang."),
    ("cartApi.updateItem", "PUT /cart/items/{id}", "CartController.UpdateItem", "OrderService.UpdateCartItemAsync", "CHITIET_GIOHANG, BIENSANPHAM, SANPHAM", "Cap nhat so luong, validate ton kho."),
    ("cartApi.removeItem", "DELETE /cart/items/{id}", "CartController.RemoveItem", "OrderService.RemoveCartItemAsync", "CHITIET_GIOHANG", "Xoa mot dong gio hang cua user."),
    ("cartApi.clear", "DELETE /cart/clear", "CartController.Clear", "OrderService.ClearCartAsync", "CHITIET_GIOHANG, GIOHANG", "Xoa toan bo dong trong gio hang active."),
    ("orderApi.getAll", "GET /orders", "OrdersController.GetOrders", "OrderRepository.SearchAsync", "DONHANG, CHITIET_DONHANG, THANHTOAN, HOSO_TRAGOP, YEUCAU_HOANTIEN", "Lay don hang cua user hoac theo filter trang thai."),
    ("orderApi.getById", "GET /orders/{id}", "OrdersController.GetOrderById", "OrderRepository.GetByIdAsync", "DONHANG, CHITIET_DONHANG, DONHANG_LICHSU_TRANGTHAI, THANHTOAN, DONHANG_VOUCHER, HOSO_TRAGOP, YEUCAU_HOANTIEN", "Lay chi tiet day du don, dong hang, lich su, thanh toan, voucher, tra gop, hoan tien."),
    ("orderApi.create", "POST /orders", "OrdersController.CreateFromCart", "OrderService.CreateFromCartAsync + OrderRepository", "GIOHANG, CHITIET_GIOHANG, DONHANG, CHITIET_DONHANG, DONHANG_LICHSU_TRANGTHAI, TONKHO_GIUCHO, THANHTOAN, HOSO_TRAGOP, DONHANG_VOUCHER, VOUCHER_*", "Tao don tu gio hang, validate voucher, giu ton, tao thanh toan/dat coc/tra gop, ghi nhan voucher."),
    ("orderApi.getShippingQuote", "POST /orders/shipping-quote", "OrdersController.GetShippingQuote", "OrderService.GetShippingQuoteAsync", "VOUCHER, VOUCHER_NGUOIDUNG, GIOHANG", "Tinh phi giao hang va giam phi neu voucher hop le."),
    ("orderApi.getPaymentInfo", "GET /orders/{id}/payment-info", "OrdersController.GetPaymentInfo", "OrderService.GetPaymentInfoAsync", "DONHANG, THANHTOAN, HETHONG_CAUHINH", "Lay so tien can thanh toan, noi dung chuyen khoan, cau hinh ngan hang."),
    ("orderApi.cancel", "PUT /orders/{id}/cancel", "OrdersController.CancelOrder", "OrderService.CancelOrderAsync", "DONHANG, DONHANG_LICHSU_TRANGTHAI, TONKHO_GIUCHO, VOUCHER_*", "Cap nhat don Cancelled, giai phong ton giu cho, huy ghi nhan voucher."),
    ("orderApi.requestRefund", "POST /orders/{id}/request-refund", "OrdersController.RequestRefund", "OrderService.RequestRefundAsync", "YEUCAU_HOANTIEN, DONHANG", "Tao yeu cau hoan tien cho don da thanh toan/du dieu kien."),
    ("paymentApi.getByOrder", "GET /payments/order/{orderId}", "PaymentsController.GetByOrder", "PaymentRepository.GetByOrderIdAsync", "THANHTOAN, DONHANG", "Lay danh sach giao dich thanh toan cua don."),
    ("voucherApi.getAll", "GET /vouchers/active", "VouchersController.GetActive", "OrderDbContext LINQ", "VOUCHER", "Lay voucher dang hoat dong, con han, con luot va public."),
    ("voucherApi.save", "POST /vouchers/save", "VouchersController.SaveVoucher", "OrderDbContext LINQ", "VOUCHER, VOUCHER_NGUOIDUNG", "Luu voucher vao tai khoan user."),
    ("voucherApi.getMine", "GET /vouchers/my", "VouchersController.GetMine", "OrderDbContext LINQ", "VOUCHER_NGUOIDUNG, VOUCHER", "Lay voucher user da luu."),
    ("voucherApi.getMineCount", "GET /vouchers/my/count", "VouchersController.GetMineCount", "OrderDbContext LINQ", "VOUCHER_NGUOIDUNG, VOUCHER", "Dem voucher da luu con kha dung."),
    ("voucherApi.getApplicable", "POST /vouchers/applicable", "VouchersController.GetApplicable", "OrderDbContext + raw SQL", "VOUCHER, VOUCHER_SANPHAM, VOUCHER_DANHMUC, VOUCHER_HANGXE", "Loc voucher ap dung theo san pham/danh muc/hang xe trong gio/checkout."),
    ("voucherApi.validate", "POST /vouchers/validate", "VouchersController.ValidateVoucher", "sp_Voucher_KiemTraTruocKhiTaoDon", "GIOHANG, CHITIET_GIOHANG, VOUCHER, VOUCHER_*", "Kiem tra voucher truoc khi tao don, tinh so tien giam."),
    ("userApi.getProfile", "GET /users/me", "UsersController.GetMe", "AuthDbContext", "NGUOIDUNG, VAITRO, NGUOIDUNG_VAITRO", "Lay profile user hien tai."),
    ("userApi.updateProfile", "PUT /users/me", "UsersController.UpdateMe", "AuthDbContext", "NGUOIDUNG", "Cap nhat ho ten, email, so dien thoai."),
    ("userApi.changePassword", "PUT /users/me/password", "UsersController.ChangePassword", "AuthDbContext", "NGUOIDUNG", "Kiem tra mat khau cu, cap nhat hash mat khau moi."),
    ("userApi.getAddresses", "GET /users/me/addresses", "UsersController.GetMyAddresses", "AuthDbContext", "NGUOIDUNG_DIACHI", "Lay danh sach dia chi cua user."),
    ("userApi.createAddress", "POST /users/me/addresses", "UsersController.CreateMyAddress", "AuthDbContext", "NGUOIDUNG_DIACHI", "Tao dia chi; neu la mac dinh thi bo mac dinh cac dia chi khac."),
    ("userApi.updateAddressById", "PUT /users/me/addresses/{id}", "UsersController.UpdateMyAddress", "AuthDbContext", "NGUOIDUNG_DIACHI", "Sua dia chi thuoc user."),
    ("userApi.setDefaultAddress", "PUT /users/me/addresses/{id}/default", "UsersController.SetDefaultAddress", "AuthDbContext", "NGUOIDUNG_DIACHI", "Dat dia chi mac dinh, clear flag o cac dia chi khac."),
    ("userApi.deleteAddress", "DELETE /users/me/addresses/{id}", "UsersController.DeleteMyAddress", "AuthDbContext", "NGUOIDUNG_DIACHI", "Xoa dia chi thuoc user."),
    ("favoriteApi.getMine", "GET /favorites", "FavoritesController.GetMine", "CatalogDbContext", "YEUTHICH, SANPHAM", "Lay danh sach yeu thich cua user, kem thong tin san pham."),
    ("favoriteApi.add", "POST /favorites/{productId}", "FavoritesController.Add", "CatalogDbContext", "YEUTHICH, SANPHAM", "Kiem tra san pham ton tai va insert neu chua co."),
    ("favoriteApi.remove", "DELETE /favorites/{productId}", "FavoritesController.Remove", "CatalogDbContext", "YEUTHICH", "Xoa dong yeu thich theo user + san pham."),
    ("contentApi.getFaqs", "GET /content/faqs", "ContentController.GetFaqs", "CatalogDbContext", "FAQ", "Lay FAQ dang hien thi cho storefront."),
    ("contentApi.createContactRequest", "POST /content/contact-requests", "ContentController.CreateContactRequest", "CatalogDbContext", "LIENHE_YEUCAU, SANPHAM", "Tao yeu cau lien he, optional lien ket san pham."),
]


SERVICE_TABLES = [
    ("AuthService", "AuthDbContext", "NGUOIDUNG, NGUOIDUNG_DIACHI, VAITRO, NGUOIDUNG_VAITRO, MATKHAU_DATLAI", "Dang nhap, dang ky, ho so, dia chi, reset password."),
    ("CatalogService", "CatalogDbContext", "SANPHAM, BIENSANPHAM, ANHSANPHAM, DANHMUC, HANGXE, DONGXE, DANHGIASANPHAM, YEUTHICH, FAQ, LIENHE_YEUCAU, BAIVIET, BANNER_TRANGCHU, DONHANG, CHITIET_DONHANG", "Catalog storefront, danh gia, yeu thich, FAQ/lien he. DONHANG/CHITIET_DONHANG duoc doc de xac minh quyen danh gia."),
    ("OrderService", "OrderDbContext", "GIOHANG, CHITIET_GIOHANG, DONHANG, CHITIET_DONHANG, DONHANG_LICHSU_TRANGTHAI, TONKHO_GIUCHO, THANHTOAN, HOSO_TRAGOP, YEUCAU_HOANTIEN, VOUCHER, VOUCHER_NGUOIDUNG, DONHANG_VOUCHER, VOUCHER_SANPHAM, VOUCHER_DANHMUC, VOUCHER_HANGXE, HETHONG_CAUHINH", "Gio hang, checkout, don hang, voucher, giu ton, thanh toan/dat coc/tra gop, hoan tien."),
    ("PaymentService", "PaymentDbContext", "THANHTOAN, DONHANG, TONKHO_GIUCHO, SANPHAM, BIENSANPHAM, NGUOIDUNG", "Frontend khach hien chi doc thanh toan theo don; admin/confirm payment khong nam trong scope Frontend."),
]


STORED_PROCS = [
    ("sp_Voucher_KiemTraTruocKhiTaoDon", "Voucher validate va tao don", "Kiem tra voucher theo user/gio hang/phi van chuyen, tra valid/message/discount."),
    ("sp_Voucher_GhiNhanSuDung", "Tao don thanh cong", "Ghi nhan voucher da dung vao DONHANG_VOUCHER/VOUCHER_NGUOIDUNG va tang so lan da dung."),
    ("sp_Voucher_HuySuDungTheoDon", "Huy don", "Rollback usage voucher khi don bi huy."),
    ("sp_TONKHO_ApDungBienDong", "Tao/huy/confirm don va mot so nghiep vu ton kho", "Cap nhat ton/giu cho theo MaSanPham/MaBienSanPham va ghi log ton kho."),
]


SEQUENCES = [
    ("Dang nhap", [
        "LoginPage submit email/password.",
        "authApi.login POST /auth/login qua API Gateway.",
        "AuthController.Login goi AuthService kiem tra NGUOIDUNG va role.",
        "JWT tra ve frontend, httpClient gan Authorization bearer cho request tiep theo.",
    ]),
    ("Xem danh sach san pham", [
        "ProductListPage tao query tu filter UI.",
        "productApi.getFilters doc DANHMUC/HANGXE/DONGXE; productApi.getAll doc SANPHAM.",
        "CatalogService map anh chinh, brand/category name, review summary.",
        "ProductCard render gia, anh, danh gia, link chi tiet.",
    ]),
    ("Them vao gio hang", [
        "Nguoi dung chon san pham/bien the va so luong.",
        "Frontend goi productApi.getById neu quick-add can kiem tra bien the.",
        "cartApi.addItem POST /cart/items voi MaSanPham, MaBienSanPham, SoLuong.",
        "OrderService tao/lay GIOHANG, validate SANPHAM/BIENSANPHAM va upsert CHITIET_GIOHANG.",
    ]),
    ("Checkout tao don", [
        "CheckoutPage lay dia chi user va gio hang hien tai.",
        "voucherApi.getApplicable/validate va orderApi.getShippingQuote tinh voucher/phi van chuyen.",
        "orderApi.create POST /orders gui thong tin nhan hang, payment method, voucher, dat coc/tra gop.",
        "OrderService tao DONHANG/CHITIET_DONHANG, TONKHO_GIUCHO, THANHTOAN/HOSO_TRAGOP, lich su trang thai; goi stored procedure voucher/ton kho khi can.",
    ]),
    ("Huy don", [
        "PaymentPage/OrderDetailPage goi orderApi.cancel PUT /orders/{id}/cancel.",
        "OrderService validate quyen va trang thai don.",
        "Cap nhat DONHANG + DONHANG_LICHSU_TRANGTHAI, giai phong TONKHO_GIUCHO, huy voucher da ghi nhan.",
    ]),
    ("Danh gia san pham", [
        "ProductReviews/OrdersPage goi reviewApi.getMine de biet co the danh gia.",
        "Backend doc DONHANG/CHITIET_DONHANG de xac minh user da mua va don da giao/hoan tat.",
        "reviewApi.create POST multipart /products/{id}/reviews.",
        "CatalogService luu DANHGIASANPHAM trang thai Pending, summary chi tinh Approved.",
    ]),
]


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    p = doc.add_paragraph(style="Title")
    p.add_run("Luồng chức năng Frontend -> Backend -> Database")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Phạm vi: storefront tại Frontend/src, không bao gồm FrontendAdmin. Gateway mặc định: /api -> http://localhost:5000.")

    add_note(
        doc,
        "Cách đọc",
        "Mỗi dòng đi từ màn hình hoặc helper frontend, qua api.js/httpClient, tới endpoint backend và các bảng hoặc stored procedure chính. Danh sách tập trung vào luồng đang được Frontend gọi thực tế.",
    )

    add_heading(doc, "1. Tổng quan kiến trúc", 1)
    add_bullets(doc, [
        "Frontend dùng React Router trong Frontend/src/App.jsx; các route bảo vệ đi qua ProtectedRoute và token JWT.",
        "Frontend/src/services/httpClient.js tạo axios baseURL mặc định /api, dev proxy chuyển tới API Gateway localhost:5000.",
        "API Gateway Ocelot route /api/auth và /api/users sang AuthService:5001; /api/products, /api/categories, /api/content, /api/favorites, /api/reviews sang CatalogService:5002; /api/cart, /api/orders, /api/vouchers sang OrderService:5003; /api/payments sang PaymentService:5004.",
        "Frontend/src/services/api.js là lớp contract chính: map payload tiếng Anh ở UI sang field tiếng Việt/backend và normalize response về shape ổn định.",
    ])

    add_heading(doc, "2. Route/page storefront và API được gọi", 1)
    add_table(
        doc,
        ["Route", "Frontend entry", "Chức năng", "API helper", "Endpoint", "Service"],
        PAGE_ROUTES,
        [1050, 1700, 2350, 1700, 1700, 860],
        font_size=7.5,
    )

    add_heading(doc, "3. Ma trận endpoint -> backend -> database", 1)
    add_table(
        doc,
        ["Frontend API", "Endpoint", "Controller/action", "Backend layer", "Bảng/SP chính", "Truy vấn/thao tác"],
        ENDPOINTS,
        [1260, 1500, 1600, 1500, 2000, 1500],
        font_size=7.2,
    )

    add_heading(doc, "4. Luồng nghiệp vụ chính theo thứ tự xử lý", 1)
    for title, steps in SEQUENCES:
        add_heading(doc, title, 2)
        add_numbered(doc, steps)

    add_heading(doc, "5. Bảng dữ liệu theo service", 1)
    add_table(
        doc,
        ["Service", "DbContext", "Bảng chính liên quan Frontend", "Vai trò trong storefront"],
        SERVICE_TABLES,
        [1200, 1500, 4200, 2460],
        font_size=8.0,
    )

    add_heading(doc, "6. Stored procedure/raw SQL có ảnh hưởng trực tiếp", 1)
    add_table(
        doc,
        ["Stored procedure / raw SQL", "Luồng gọi", "Mục đích"],
        STORED_PROCS,
        [2600, 2200, 4560],
        font_size=8.2,
    )

    add_heading(doc, "7. Ghi chú về ranh giới dữ liệu", 1)
    add_bullets(doc, [
        "CatalogService và OrderService cùng map một số bảng catalog/order để phục vụ đọc nhanh theo bounded context; tài liệu ghi service đang xử lý endpoint chứ không khẳng định ownership vật lý duy nhất của bảng.",
        "Các endpoint admin như quản lý sản phẩm, tồn kho, bảo hành, mua hàng, POS, audit log không nằm trong phạm vi vì không được Frontend khách hàng gọi.",
        "Một số bảng liên kết voucher (VOUCHER_SANPHAM, VOUCHER_DANHMUC, VOUCHER_HANGXE) không khai báo DbSet chính nhưng được truy vấn bằng raw SQL trong VouchersController.",
        "Thanh toán online thực tế có thể mở rộng bằng callback/confirm ở PaymentService, nhưng Frontend khách hàng hiện chỉ đọc trạng thái thanh toán theo đơn và thông tin chuyển khoản từ OrderService.",
    ])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("ShowRoomDB storefront flow map").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
