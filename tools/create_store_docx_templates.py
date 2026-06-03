from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "templates"
OUT.mkdir(parents=True, exist_ok=True)
BLUE = "1F4D78"
LIGHT = "E8EEF5"
FIELD = "................................"
DATE_FIELD = "....../....../.........."

def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)

def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_run(run, size=10.5, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def setup(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("MoToSale | Biểu mẫu vận hành cửa hàng"), 9, False, "666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("MoToSale - Lưu cùng hồ sơ nghiệp vụ"), 8.5, False, "777777")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run("MOTOSALE"), 12, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title.upper()), 17, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(f"Mã phiếu: {FIELD}     Ngày: {DATE_FIELD}"), 10, False, "555555")

def add_section(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label), 11, True, BLUE)

def add_fields(doc, fields, widths=(1800, 2880, 1800, 2880)):
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for left_label, left_value, right_label, right_value in fields:
        cells = table.add_row().cells
        for index, text in enumerate((left_label, left_value, right_label, right_value)):
            cells[index].text = text
            for para in cells[index].paragraphs:
                para.paragraph_format.space_after = Pt(1)
                for run in para.runs:
                    set_run(run, 9.7, index in (0, 2))
            if index in (0, 2):
                set_cell_fill(cells[index], LIGHT)
    return table

def add_items(doc, headers, widths, blank_rows=4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        set_cell_fill(table.cell(0, index), BLUE)
        for run in table.cell(0, index).paragraphs[0].runs:
            set_run(run, 9.2, True, "FFFFFF")
    for _ in range(blank_rows):
        cells = table.add_row().cells
        for cell in cells:
            cell.text = "\n"
    return table

def add_signatures(doc, labels):
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=len(labels))
    table.autofit = False
    widths = [int(9360 / len(labels))] * len(labels)
    set_table_geometry(table, widths)
    for cell, label in zip(table.rows[0].cells, labels):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(label), 10, True)
        p = cell.add_paragraph("\n\n\n(Ký và ghi rõ họ tên)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run(run, 9, False, "666666")

def repair_form():
    doc = Document()
    setup(doc, "Biên bản tiếp nhận sửa chữa")
    add_section(doc, "1. Thông tin khách hàng và xe")
    add_fields(doc, [
        ("Khách hàng", FIELD, "Điện thoại", FIELD),
        ("Biển số / số khung", FIELD, "Dòng xe", FIELD),
        ("Số km hiện tại", FIELD, "Kỹ thuật viên", FIELD),
    ])
    add_section(doc, "2. Tình trạng tiếp nhận")
    add_fields(doc, [
        ("Lỗi khách mô tả", FIELD, "Ngày hẹn trả", FIELD),
        ("Ghi chú ngoại quan", FIELD, "Trạng thái", "Tiếp nhận / Đang sửa / Hoàn thành"),
    ])
    add_section(doc, "3. Hạng mục sửa chữa và phụ tùng")
    add_items(doc, ["STT", "Hạng mục / phụ tùng", "SKU", "SL", "Đơn giá", "Thành tiền"], [520, 3430, 1450, 620, 1500, 1840], 4)
    add_section(doc, "4. Xác nhận")
    add_fields(doc, [("Tổng dự kiến", FIELD, "Tổng thực tế", FIELD)])
    add_signatures(doc, ["Khách hàng", "Kỹ thuật viên", "Người tiếp nhận"])
    doc.save(OUT / "Bien-ban-tiep-nhan-sua-chua.docx")

def warranty_form():
    doc = Document()
    setup(doc, "Phiếu tiếp nhận bảo hành")
    add_section(doc, "1. Thông tin khách hàng và sản phẩm")
    add_fields(doc, [
        ("Khách hàng", FIELD, "Điện thoại", FIELD),
        ("Sản phẩm / xe", FIELD, "Mã đơn hàng", FIELD),
        ("Số khung / serial", FIELD, "Ngày mua", DATE_FIELD),
    ])
    add_section(doc, "2. Phạm vi và tình trạng bảo hành")
    add_fields(doc, [
        ("Ngày tiếp nhận", DATE_FIELD, "Hạn bảo hành", DATE_FIELD),
        ("Lỗi ghi nhận", FIELD, "Tình trạng", "Tiếp nhận / Kiểm tra / Hoàn thành"),
        ("Kết luận", FIELD, "Chi phí ngoài BH", FIELD),
    ])
    add_section(doc, "3. Lịch sử xử lý")
    add_items(doc, ["Ngày", "Trạng thái", "Nội dung xử lý", "Người thực hiện"], [1500, 1800, 4160, 1900], 5)
    add_section(doc, "4. Xác nhận")
    add_signatures(doc, ["Khách hàng", "Kỹ thuật viên", "Đại diện cửa hàng"])
    doc.save(OUT / "Phieu-tiep-nhan-bao-hanh.docx")

if __name__ == "__main__":
    repair_form()
    warranty_form()
