from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_frontend_flow_docx import (
    API_CONTRACTS,
    BACKEND_CODE_MAP,
    DB_TABLE_DETAILS,
    ENDPOINTS,
    FRONTEND_CODE_MAP,
    OUT as BASE_OUT,
    PAGE_CODE_DETAILS,
    PAGE_ROUTES,
    SEQUENCES,
    SERVICE_TABLES,
    STORED_PROCS,
    add_bullets,
    add_code_block,
    add_heading,
    add_note,
    add_numbered,
    add_table,
    setup_styles,
)


OUT = BASE_OUT.with_name("frontend-backend-database-flows-detailed.docx")


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Luong chuc nang Frontend -> Backend -> Database")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Ban chi tiet: giai thich code, API contract, backend layer va database. "
        "Pham vi chi tinh storefront Frontend/src, khong bao gom FrontendAdmin."
    )

    add_note(
        doc,
        "Cach doc",
        "Tai lieu nay doc theo dung duong di cua request: Page/Component -> Context/API helper -> "
        "httpClient/proxy/gateway -> Controller -> Service/Repository -> DbContext/raw SQL/stored procedure. "
        "Cac bang ghi trong tai lieu la bang chinh bi doc/ghi trong luong, khong phai toan bo schema.",
    )

    add_heading(doc, "1. Tong quan kien truc", 1)
    add_bullets(doc, [
        "Frontend chay tren Vite port 5174. Code khach hang nam trong Frontend/src.",
        "Frontend khong goi truc tiep service 5001/5002/5003/5004; moi request dung baseURL /api va duoc proxy toi API Gateway localhost:5000.",
        "API Gateway Ocelot route theo prefix: auth/users -> AuthService, products/categories/content/favorites/reviews -> CatalogService, cart/orders/vouchers -> OrderService, payments -> PaymentService.",
        "httpClient.js la lop gan token va xu ly error message. api.js la lop contract nghiep vu. normalizers.js/productMappers.js bien response nhieu kieu key thanh model on dinh cho UI.",
        "Backend di theo pattern controller -> service/repository -> DbContext. Mot so luong quan trong dung raw SQL/stored procedure: voucher validation, voucher usage, inventory movement.",
    ])

    add_code_block(doc, "Request pipeline tong quat", [
        "React page/component",
        "  -> Context (Auth/Cart/Favorite) hoac services/api.js",
        "  -> normalizers.js / productMappers.js map query, body, response",
        "  -> httpClient.js axios instance",
        "  -> Authorization: Bearer <JWT> neu user da dang nhap",
        "  -> Vite proxy /api -> API Gateway localhost:5000",
        "  -> Ocelot chon downstream service",
        "  -> ASP.NET Controller",
        "  -> Service/Repository",
        "  -> EF Core DbContext, raw SQL, hoac stored procedure",
    ])

    add_heading(doc, "2. Giai thich code Frontend theo lop", 1)
    add_table(
        doc,
        ["File", "Vai tro", "Giai thich code"],
        FRONTEND_CODE_MAP,
        [2300, 1700, 5360],
        font_size=8.0,
    )

    add_heading(doc, "3. Giai thich code theo man hinh", 1)
    add_table(
        doc,
        ["Frontend file", "Code chinh dang lam gi", "Y nghia voi backend/database"],
        PAGE_CODE_DETAILS,
        [2100, 3600, 3660],
        font_size=7.7,
    )

    add_heading(doc, "4. Contract API trong Frontend/src/services/api.js", 1)
    add_table(
        doc,
        ["API helper", "Input tu UI", "Request gui backend", "Response/side effect tren UI"],
        API_CONTRACTS,
        [1500, 2300, 3100, 2460],
        font_size=7.6,
    )

    add_heading(doc, "5. Route/page storefront va API duoc goi", 1)
    add_table(
        doc,
        ["Route", "Frontend entry", "Chuc nang", "API helper", "Endpoint", "Service"],
        PAGE_ROUTES,
        [1050, 1700, 2350, 1700, 1700, 860],
        font_size=7.5,
    )

    add_heading(doc, "6. Giai thich code Backend theo service", 1)
    add_table(
        doc,
        ["Backend file", "Vai tro", "Giai thich code"],
        BACKEND_CODE_MAP,
        [2700, 1800, 4860],
        font_size=7.8,
    )

    add_heading(doc, "7. Ma tran endpoint -> backend -> database", 1)
    add_table(
        doc,
        ["Frontend API", "Endpoint", "Controller/action", "Backend layer", "Bang/SP chinh", "Truy van/thao tac"],
        ENDPOINTS,
        [1260, 1500, 1600, 1500, 2000, 1500],
        font_size=7.2,
    )

    add_heading(doc, "8. Luong nghiep vu chinh theo thu tu xu ly", 1)
    for heading, steps in SEQUENCES:
        add_heading(doc, heading, 2)
        add_numbered(doc, steps)

    add_code_block(doc, "Pseudo-code: checkout tao don", [
        "CheckoutPage submit",
        "  subtotal = sum(cart.items.lineTotal)",
        "  optional voucherApi.validate(code, subtotal, productIds, categoryIds, brandIds, shippingFee)",
        "  quote = orderApi.getShippingQuote(receivingMethod, province, voucherCode, orderType)",
        "  payload = shipping fields + voucherCode + paymentMethod + installment/deposit",
        "  POST /orders",
        "OrderService.CreateFromCart",
        "  load active cart + cart items",
        "  validate products, variants, stock, voucher",
        "  recompute subtotal, discount, shipping on server",
        "  create DONHANG + CHITIET_DONHANG + DONHANG_LICHSU_TRANGTHAI",
        "  create TONKHO_GIUCHO or apply inventory movement depending on flow",
        "  create THANHTOAN and optional HOSO_TRAGOP",
        "  record voucher usage via stored procedure",
    ])

    add_code_block(doc, "Pseudo-code: them vao gio hang", [
        "ProductCard/ProductDetailPage",
        "  if product has variants and no selected variant: navigate detail",
        "  payload = { productId, variantId, quantity }",
        "cartApi.addItem",
        "  POST /cart/items { maSanPham, maBienSanPham, soLuong }",
        "OrderService.AddCartItemAsync",
        "  get or create active GIOHANG",
        "  validate SANPHAM/BIENSANPHAM status and stock",
        "  insert/update CHITIET_GIOHANG",
        "  return normalized cart with product/variant/price/image snapshot",
    ])

    add_heading(doc, "9. Bang du lieu theo service", 1)
    add_table(
        doc,
        ["Service", "DbContext", "Bang chinh lien quan Frontend", "Vai tro trong storefront"],
        SERVICE_TABLES,
        [1200, 1500, 4200, 2460],
        font_size=8.0,
    )

    add_heading(doc, "10. Giai thich bang database", 1)
    add_table(
        doc,
        ["Bang", "Service doc/ghi", "Y nghia trong luong Frontend"],
        DB_TABLE_DETAILS,
        [2600, 2100, 4660],
        font_size=7.8,
    )

    add_heading(doc, "11. Stored procedure/raw SQL co anh huong truc tiep", 1)
    add_table(
        doc,
        ["Stored procedure / raw SQL", "Luong goi", "Muc dich"],
        STORED_PROCS,
        [2600, 2200, 4560],
        font_size=8.2,
    )

    add_heading(doc, "12. Diem can chu y khi doc/sua code", 1)
    add_bullets(doc, [
        "Neu UI hien sai du lieu, kiem tra theo thu tu: component state -> api.js mapping -> normalizers.js/productMappers.js -> response backend -> query database.",
        "Khong sua component de goi endpoint truc tiep neu api.js da co helper; nen giu contract tap trung trong services/api.js.",
        "Khong tin gia/tong tien tinh o frontend khi checkout. Frontend chi preview; backend OrderService phai la noi tinh va snapshot don hang cuoi.",
        "Khi them field response backend, cap nhat mapper alias o normalizers.js hoac productMappers.js de UI khong phu thuoc casing/tieng Viet/tieng Anh.",
        "CatalogService va OrderService co the cung map SANPHAM/BIENSANPHAM vi bounded context khac nhau; can xem endpoint dang di vao service nao truoc khi sua query.",
        "FrontendAdmin khong nam trong tai lieu nay. Cac endpoint admin trong controller co the xuat hien cung file nhung khong duoc liet ke la luong storefront.",
    ])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("ShowRoomDB storefront detailed flow map").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
