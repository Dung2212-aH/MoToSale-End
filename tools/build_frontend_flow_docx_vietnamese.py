from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_frontend_flow_docx import (
    add_bullets,
    add_code_block,
    add_heading,
    add_note,
    add_numbered,
    add_table,
    setup_styles,
)

from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "frontend-backend-database-flows-vietnamese-detailed.docx"


FRONTEND_CODE = [
    (
        "Frontend/src/App.jsx",
        "File khởi tạo toàn bộ ứng dụng khách hàng.",
        "Code dựng BrowserRouter, bọc các Provider theo thứ tự AuthProvider -> NotificationProvider -> FavoriteProvider -> CartProvider, rồi khai báo Routes. Ý nghĩa kỹ thuật: AuthProvider phải nằm ngoài Cart/Favorite để các context này biết user đã đăng nhập hay chưa. Route công khai như /products được render trực tiếp; route cần đăng nhập như /cart, /checkout, /orders, /account được bọc bằng ProtectedRoute.",
    ),
    (
        "Frontend/src/components/ProtectedRoute.jsx",
        "Chặn truy cập trang cần đăng nhập.",
        "Code đọc isAuthenticated và loading từ AuthContext. Khi AuthContext còn khởi tạo thì hiển thị loading. Nếu chưa đăng nhập thì tạo redirect bằng pathname/search/hash hiện tại rồi điều hướng sang /login?redirect=... Nếu đã đăng nhập thì trả children. Chức năng code không gọi API, nhưng quyết định luồng request sau đó có được phép chạy hay không.",
    ),
    (
        "Frontend/src/services/httpClient.js",
        "Lớp hạ tầng HTTP dùng chung.",
        "Code tạo axios instance với baseURL mặc định là /api. File này quản lý token trong sessionStorage/localStorage, decode JWT để lấy claim, kiểm tra token hết hạn, lưu/xóa user, phát AUTH_CHANGED_EVENT, và gắn interceptor. Request interceptor tự thêm Authorization: Bearer <token>; response interceptor lấy message từ backend và gán vào error.message để UI hiển thị lỗi đúng.",
    ),
    (
        "Frontend/src/services/api.js",
        "Lớp hợp đồng API nghiệp vụ.",
        "Code gom các nhóm authApi, productApi, cartApi, orderApi, voucherApi, userApi, favoriteApi, contentApi. Mỗi hàm nhận dữ liệu theo tên thân thiện ở UI, đổi sang payload backend đang dùng, gọi endpoint, rồi map response về shape ổn định. Đây là nơi quan trọng nhất nếu cần sửa payload hoặc endpoint cho Frontend khách hàng.",
    ),
    (
        "Frontend/src/services/normalizers.js",
        "Chuẩn hóa dữ liệu trả về và query gửi đi.",
        "Code có field() để đọc nhiều tên khóa khác nhau như id/Id/maSanPham/MaSanPham; toQuery() đổi categoryId thành MaDanhMuc, brandId thành MaHangXe, minPrice thành GiaTu, sortBy thành SortBy/SortDescending; mapOrder/mapPayment/mapVoucher/mapFavorite/mapReview gom response backend về object frontend dùng được.",
    ),
    (
        "Frontend/src/utils/productMappers.js",
        "Chuẩn hóa sản phẩm, danh mục, giỏ hàng.",
        "normalizeProduct nhận raw product từ backend rồi tạo id, name, basePrice, salePrice, discountPercent, images, variants. normalizeProductList xử lý response phân trang. normalizeCart map GIOHANG/CHITIET_GIOHANG thành cart.items có unitPrice, lineTotal, product, productVariant. Code này quyết định ProductCard, ProductDetail, Cart hiển thị field nào.",
    ),
    (
        "Frontend/src/utils/productOptions.js",
        "Tạo lựa chọn phiên bản/màu cho trang chi tiết.",
        "Code không gọi backend. Nó đọc product.variants và product.images đã normalize, suy ra danh sách version, color, ảnh theo biến thể, ảnh theo màu, selectedVariant, fallback nếu backend thiếu mapping màu. Đây là lớp biến dữ liệu API thành trạng thái UI có thể chọn.",
    ),
    (
        "Frontend/src/contexts/AuthContext.jsx",
        "Quản lý phiên đăng nhập trong React.",
        "Code khởi tạo user từ authApi.getCurrentUser(), expose login/register/logout/updateUser/isAuthenticated. Khi login thành công, authApi lưu token rồi AuthContext cập nhật user. Context lắng nghe storage và AUTH_CHANGED_EVENT để các tab hoặc context khác đồng bộ trạng thái đăng nhập.",
    ),
    (
        "Frontend/src/contexts/CartContext.jsx",
        "Quản lý giỏ hàng dùng chung toàn ứng dụng.",
        "Khi isAuthenticated thay đổi, code gọi refreshCart() -> cartApi.getMine(). Các hàm addItem/updateItem/removeItem/clearCart gọi backend rồi applyCart để cập nhật cart và count. updateItem còn kiểm tra tồn kho local trước khi gửi PUT. Context này là nguồn dữ liệu cho Header mini-cart, CartPage và CheckoutPage.",
    ),
    (
        "Frontend/src/contexts/FavoriteContext.jsx",
        "Quản lý danh sách yêu thích.",
        "Code load favoriteApi.getMine() khi user đã đăng nhập, tạo Set favoriteIds để kiểm tra nhanh isFavorite. toggleFavorite dùng optimistic update: cập nhật UI trước, gọi POST/DELETE backend sau; nếu backend lỗi thì rollback về danh sách cũ.",
    ),
    (
        "Frontend/vite.config.js",
        "Cấu hình dev server và proxy.",
        "Frontend chạy port 5174. Proxy /api và /uploads sang http://localhost:5000. Nhờ vậy code frontend chỉ gọi /api/products, /api/orders..., còn việc route sang service nào do API Gateway quyết định.",
    ),
]


PAGE_CODE = [
    (
        "HomePage.jsx + Header.jsx",
        "HomePage dùng useEffect để gọi productApi.getAll và categoryApi.getAll; nếu user đã đăng nhập thì gọi thêm orderApi.getMyOrders. Header gọi productApi.getFilters để lấy dữ liệu tìm kiếm/lọc và voucherApi.getMineCount để hiện số voucher đã lưu.",
        "Code chỉ đọc dữ liệu. Dữ liệu sản phẩm đi qua normalizeProductList trước khi render ProductCard. Nếu bấm thêm nhanh mà sản phẩm có biến thể, code gọi productApi.getById để kiểm tra rồi điều hướng sang trang chi tiết.",
    ),
    (
        "ProductListPage.jsx",
        "Code giữ state filter, sort, page; build apiQueryValues; gọi song song productApi.getFilters và productApi.getAll. Khi người dùng đổi filter, query UI được chuyển thành query backend qua toQuery().",
        "categoryId/brandId/minPrice/maxPrice ở UI không gửi nguyên dạng xuống backend; api.js/normalizers.js đổi thành MaDanhMuc/MaHangXe/GiaTu/GiaDen. Backend ProductRepository mới là nơi apply filter thật trên SANPHAM.",
    ),
    (
        "ProductDetailPage.jsx",
        "Code lấy id từ URL, gọi productApi.getById(id), tạo options bằng normalizeProductOptions(product), tính selectedVariant theo selectedVersion/selectedColor, xử lý chọn ảnh, chọn màu, chọn phiên bản, thêm giỏ hàng, mua ngay, yêu thích và sản phẩm liên quan.",
        "Trang này là nơi ghép nhiều bounded context: CatalogService cho sản phẩm/review/favorite, OrderService cho cart. buildCartPayload() chỉ gửi productId, variantId, quantity; giá và tồn kho cuối cùng do backend cart/order xác nhận.",
    ),
    (
        "ProductReviews.jsx / ReviewModal.jsx",
        "Code load reviewApi.getByProduct, reviewApi.getSummary và reviewApi.getMine. Khi submit, buildReviewForm tạo multipart/form-data gồm Diem, NoiDung, TieuDe, Image, MaDonHang nếu có.",
        "Backend chỉ public review Approved. getMine trả canReview/hasPurchased/eligibleOrderId để UI biết có được mở form đánh giá không. Review mới được lưu Pending nên không chắc xuất hiện ngay trong danh sách public.",
    ),
    (
        "CartPage.jsx + CartContext.jsx",
        "CartPage không tự gọi axios trực tiếp mà dùng CartContext. Context gọi GET /cart khi đăng nhập, mutation thì gọi POST/PUT/DELETE rồi cập nhật state cart/count.",
        "Frontend có thể kiểm tra số lượng vượt tồn kho để báo sớm, nhưng backend OrderService vẫn là nơi xác nhận tồn kho/gia thật. Sau update, frontend lấy lại cart server để tránh lệch tổng tiền.",
    ),
    (
        "CheckoutPage.jsx",
        "Code đọc cart từ CartContext, load userApi.getAddresses, tính subtotal preview, gọi voucherApi.getApplicable/validate, gọi orderApi.getShippingQuote, rồi build payload orderApi.create.",
        "Đây là luồng ghi database lớn nhất: backend tạo DONHANG, CHITIET_DONHANG, lịch sử trạng thái, giữ tồn, thanh toán, hồ sơ trả góp, voucher. Frontend chỉ gửi lựa chọn và thông tin nhận hàng; backend phải tính lại tiền.",
    ),
    (
        "PaymentPage.jsx",
        "Code đọc orderId, gọi orderApi.getById và orderApi.getPaymentInfo để hiển thị số tiền, nội dung chuyển khoản/QR, trạng thái thanh toán. Nếu người dùng hủy, gọi orderApi.cancel.",
        "Hủy đơn không xóa dữ liệu. Backend đổi trạng thái DONHANG, ghi DONHANG_LICHSU_TRANGTHAI, giải phóng TONKHO_GIUCHO và rollback voucher nếu đã ghi nhận.",
    ),
    (
        "OrdersPage.jsx",
        "Code gọi orderApi.getMyOrders để lấy danh sách đơn của user, có thể gọi orderApi.getAll({ trangThaiDonHang: 'AwaitingPayment' }) để phục vụ nhóm chờ thanh toán. Khi user muốn đánh giá, gọi reviewApi.getMine(productId).",
        "Dữ liệu đơn đến từ OrderService; quyền đánh giá lại được kiểm tra ở CatalogService vì DANHGIASANPHAM nằm trong catalog context và cần đối chiếu DONHANG/CHITIET_DONHANG.",
    ),
    (
        "OrderDetailPage.jsx",
        "Code load orderApi.getById, paymentApi.getByOrder và orderApi.getPaymentInfo. Handler cancel gọi PUT /orders/{id}/cancel; handler requestRefund gọi POST /orders/{id}/request-refund.",
        "Trang này đọc dữ liệu từ cả OrderService và PaymentService. Yêu cầu hoàn tiền tạo YEUCAU_HOANTIEN; payment list đọc THANHTOAN theo MaDonHang.",
    ),
    (
        "AccountPage.jsx",
        "Code gọi song song userApi.getProfile và userApi.getAddresses. Các form cập nhật profile, đổi mật khẩu, thêm/sửa/xóa/đặt mặc định địa chỉ đều đi qua userApi.",
        "Toàn bộ luồng này thuộc AuthService. Bảng chính là NGUOIDUNG và NGUOIDUNG_DIACHI. Checkout chỉ đọc địa chỉ đã lưu, còn snapshot địa chỉ giao hàng được lưu riêng vào DONHANG khi tạo đơn.",
    ),
    (
        "FavoritesPage.jsx + FavoriteContext.jsx",
        "FavoriteContext load danh sách yêu thích khi login. Page render favoriteProducts, cho bỏ yêu thích hoặc thêm nhanh vào giỏ.",
        "GET /favorites join YEUTHICH với SANPHAM để trả product nested. favoriteApi.add/remove chỉ thay đổi bảng YEUTHICH theo user hiện tại và productId.",
    ),
    (
        "VouchersPage.jsx",
        "Code gọi voucherApi.getAll để lấy voucher active public; nếu đã đăng nhập thì gọi voucherApi.getMine để biết voucher đã lưu; nút lưu gọi voucherApi.save(code).",
        "Voucher nằm trong OrderService vì dùng cho checkout/order. VOUCHER là định nghĩa mã, VOUCHER_NGUOIDUNG là quan hệ user đã lưu/đã dùng.",
    ),
    (
        "ContactPage.jsx / FaqPage.jsx",
        "ContactPage validate form rồi gọi contentApi.createContactRequest. FaqPage gọi contentApi.getFaqs và lọc/search phía client.",
        "Contact ghi LIENHE_YEUCAU, FAQ chỉ đọc FAQ active. Cả hai đi CatalogService/ContentController.",
    ),
]


API_CONTRACTS = [
    ("authApi.login", "username, password, rememberMe", "POST /auth/login body { email: username, matKhau: password }", "Lưu token/user, phát AUTH_CHANGED_EVENT, các request sau tự có Bearer token."),
    ("authApi.register", "name, email, phone, password", "POST /auth/register body { hoTen, email, soDienThoai, matKhau }", "Backend tạo NGUOIDUNG và role khách hàng."),
    ("productApi.getAll", "keyword, categoryId, brandId, minPrice, maxPrice, sortBy, page", "GET /products với DangHoatDong=true; toQuery map sang MaDanhMuc, MaHangXe, GiaTu, GiaDen, SortBy", "normalizeProductList trả items để ProductCard render."),
    ("productApi.getById", "id sản phẩm", "GET /products/{id}", "normalizeProduct trả product detail, variants, images cho ProductDetailPage."),
    ("reviewApi.create", "rating, comment, title, image, orderId", "POST multipart /products/{id}/reviews; buildReviewForm map rating->Diem, comment->NoiDung, image->Image", "Backend lưu DANHGIASANPHAM trạng thái Pending."),
    ("cartApi.addItem", "productId, variantId, quantity", "POST /cart/items body { maSanPham, maBienSanPham, soLuong }", "handleCart normalize cart và notifyCartChanged để Header/CartContext cập nhật."),
    ("cartApi.updateItem", "cart item id, quantity", "PUT /cart/items/{id} body { soLuong }", "Sau PUT gọi lại GET /cart để nhận snapshot mới."),
    ("orderApi.create", "shipping info, receivingMethod, orderType, paymentMethod, voucherCode, installment/deposit", "POST /orders body tiếng Việt: hoTenNhanHang, soDienThoaiNhanHang, diaChiNhanHang, maVoucherCode, phuongThucNhanHang, loaiDonHang, phuongThucThanhToan...", "Backend tạo đơn và trả OrderDto đã mapOrder."),
    ("orderApi.getShippingQuote", "receivingMethod, shippingProvince, voucherCode, orderType", "POST /orders/shipping-quote", "Trả shippingFee, originalShippingFee, discountAmount để checkout preview."),
    ("voucherApi.validate", "code, subtotal, productIds, categoryIds, brandIds, orderType, shippingFee", "POST /vouchers/validate", "Backend recompute theo active cart và stored procedure; trả valid/message/discountAmount."),
    ("userApi.createAddress/updateAddressById", "fullName, phoneNumber, addressLine, ward, district, province, note, isDefault", "mapAddressBody -> hoTenNhanHang, soDienThoaiNhanHang, diaChiNhanHang, ward, district, province, ghiChu, laMacDinh", "Ghi NGUOIDUNG_DIACHI."),
    ("favoriteApi.add/remove", "productId", "POST hoặc DELETE /favorites/{productId}", "FavoriteContext optimistic update và rollback nếu lỗi."),
    ("contentApi.createContactRequest", "fullName, phoneNumber, email, subject, message, inquiryType, productId", "POST /content/contact-requests body { hoTen, soDienThoai, email, tieuDe, noiDung, loaiYeuCau, maSanPham }", "Insert LIENHE_YEUCAU."),
]


ROUTES = [
    ("/", "HomePage.jsx + Header", "Tải sản phẩm, danh mục, filter header, voucher count, đơn gần đây nếu login.", "productApi.getAll; categoryApi.getAll; productApi.getFilters; orderApi.getMyOrders; voucherApi.getMineCount", "CatalogService, OrderService"),
    ("/products", "ProductListPage.jsx", "Danh sách sản phẩm, filter, sort, phân trang, quick-add.", "productApi.getFilters; productApi.getAll; productApi.getById", "CatalogService"),
    ("/products/:id", "ProductDetailPage.jsx", "Chi tiết sản phẩm, biến thể, ảnh, review, favorite, add cart.", "productApi.getById; reviewApi.*; favoriteApi.*; cartApi.addItem", "CatalogService, OrderService"),
    ("/cart", "CartPage.jsx", "Xem/sửa/xóa giỏ hàng.", "cartApi.getMine; addItem; updateItem; removeItem; clear", "OrderService"),
    ("/checkout", "CheckoutPage.jsx", "Địa chỉ, voucher, phí vận chuyển, tạo đơn.", "userApi.getAddresses; voucherApi.*; orderApi.getShippingQuote/create", "AuthService, OrderService"),
    ("/checkout/payment", "PaymentPage.jsx", "Thông tin thanh toán, hủy đơn trước thanh toán.", "orderApi.getById; getPaymentInfo; cancel", "OrderService"),
    ("/orders", "OrdersPage.jsx", "Danh sách đơn, kiểm tra quyền đánh giá.", "orderApi.getMyOrders; reviewApi.getMine", "OrderService, CatalogService"),
    ("/orders/:id", "OrderDetailPage.jsx", "Chi tiết đơn, thanh toán, hủy đơn, yêu cầu hoàn tiền.", "orderApi.getById; paymentApi.getByOrder; requestRefund", "OrderService, PaymentService"),
    ("/account", "AccountPage.jsx", "Profile, đổi mật khẩu, địa chỉ.", "userApi.getProfile/updateProfile/changePassword/address APIs", "AuthService"),
    ("/favorites", "FavoritesPage.jsx", "Danh sách yêu thích, bỏ/thêm yêu thích, quick-add.", "favoriteApi.*; productApi.getById; cartApi.addItem", "CatalogService, OrderService"),
    ("/vouchers", "VouchersPage.jsx", "Voucher active, voucher đã lưu, lưu voucher.", "voucherApi.getAll; getMine; save", "OrderService"),
    ("/contact", "ContactPage.jsx", "Gửi yêu cầu liên hệ.", "contentApi.createContactRequest", "CatalogService"),
    ("/faq", "FaqPage.jsx", "FAQ public.", "contentApi.getFaqs", "CatalogService"),
    ("/login/register/forgot-password", "Auth pages", "Đăng nhập, đăng ký, quên mật khẩu.", "authApi.login/register/forgotPassword/resetPassword", "AuthService"),
]


BACKEND_CODE = [
    ("ApiGateway/ocelot.json", "Route gateway", "Map upstream /api/... từ frontend sang service port nội bộ. Không xử lý nghiệp vụ nhưng là nơi đầu tiên cần kiểm tra nếu endpoint 404/405."),
    ("AuthController.cs", "Auth endpoints", "Nhận login/register/forgot/reset request, gọi IAuthService, trả token/user hoặc lỗi. Payload từ frontend đã được api.js map sang field backend."),
    ("UsersController.cs", "Profile và địa chỉ", "Lấy user id từ JWT claims, đọc/ghi NGUOIDUNG và NGUOIDUNG_DIACHI. Các endpoint /users/me là storefront; các endpoint admin trong cùng file không tính trong tài liệu này."),
    ("ProductsController.cs", "Catalog endpoints", "GET /products và GET /products/{id} gọi CatalogService. GET /products/filters đọc trực tiếp DANHMUC/HANGXE/DONGXE qua DbContext."),
    ("CatalogService.cs", "Map catalog DTO", "GetProductsAsync lấy product page, image map, brand/category map, review summary rồi MapProductListItem. GetProductByIdAsync lấy product, variants, images rồi MapProductDetail."),
    ("ProductRepository.cs", "Query builder sản phẩm", "Apply filter keyword/category/brand/model/type/status/price/sort/page. Đây là nơi ProductListPage tác động xuống SQL."),
    ("ReviewsController.cs", "Đánh giá", "GET public chỉ trả Approved. getMine/create kiểm tra user đã mua hàng qua DONHANG/CHITIET_DONHANG trước khi cho review."),
    ("FavoritesController.cs", "Yêu thích", "GET join YEUTHICH với SANPHAM; POST kiểm tra sản phẩm tồn tại và favorite chưa có; DELETE xóa theo user/product."),
    ("ContentController.cs", "FAQ và liên hệ", "GET /content/faqs đọc FAQ active. POST /content/contact-requests validate input và insert LIENHE_YEUCAU."),
    ("CartController.cs", "Giỏ hàng", "Controller lấy user từ JWT, gọi OrderService và trả CartDto. Frontend không tự tính authoritative total."),
    ("OrderService.cs", "Nghiệp vụ cart/order", "Xử lý cart, checkout, shipping quote, payment-info, cancel, refund. Đây là lớp có nhiều side effect database nhất trong storefront."),
    ("OrderRepository.cs", "Persistence đơn hàng", "Đọc/ghi DONHANG/CHITIET_DONHANG, giữ tồn, gọi stored procedure voucher và tồn kho khi tạo/hủy đơn."),
    ("VouchersController.cs", "Voucher", "GET active/my/count/save/applicable/validate. applicable dùng raw SQL vào bảng scope voucher; validate gọi sp_Voucher_KiemTraTruocKhiTaoDon."),
    ("PaymentsController.cs", "Thanh toán", "Frontend khách chủ yếu gọi GET /payments/order/{orderId}. Các endpoint confirm/fail/cancel phục vụ luồng khác, không phải trọng tâm storefront."),
]


ENDPOINTS = [
    ("productApi.getAll", "GET /products", "ProductsController.GetProducts -> CatalogService -> ProductRepository", "SANPHAM, ANHSANPHAM, HANGXE, DANHMUC, DANHGIASANPHAM", "Lọc/sắp xếp/phân trang sản phẩm, lấy ảnh chính, tên hãng/danh mục, điểm đánh giá."),
    ("productApi.getById", "GET /products/{id}", "ProductsController.GetProductById -> CatalogService", "SANPHAM, BIENSANPHAM, ANHSANPHAM, DANHGIASANPHAM", "Lấy chi tiết sản phẩm, biến thể, ảnh, review summary."),
    ("cartApi.getMine", "GET /cart", "CartController.GetMine -> OrderService", "GIOHANG, CHITIET_GIOHANG, SANPHAM, BIENSANPHAM, ANHSANPHAM", "Lấy giỏ hàng active của user."),
    ("cartApi.addItem", "POST /cart/items", "CartController.AddItem -> OrderService", "GIOHANG, CHITIET_GIOHANG, SANPHAM, BIENSANPHAM", "Tạo giỏ nếu chưa có, validate tồn/gia, upsert dòng giỏ."),
    ("orderApi.create", "POST /orders", "OrdersController.CreateFromCart -> OrderService -> OrderRepository", "GIOHANG, CHITIET_GIOHANG, DONHANG, CHITIET_DONHANG, TONKHO_GIUCHO, THANHTOAN, HOSO_TRAGOP, VOUCHER_*", "Tạo đơn từ giỏ, tính lại tiền, giữ tồn, tạo thanh toán/trả góp, ghi nhận voucher."),
    ("orderApi.cancel", "PUT /orders/{id}/cancel", "OrdersController.CancelOrder -> OrderService", "DONHANG, DONHANG_LICHSU_TRANGTHAI, TONKHO_GIUCHO, VOUCHER_*", "Hủy đơn, ghi lịch sử, giải phóng tồn giữ chỗ, rollback voucher."),
    ("voucherApi.validate", "POST /vouchers/validate", "VouchersController.ValidateVoucher", "GIOHANG, CHITIET_GIOHANG, VOUCHER, VOUCHER_*; sp_Voucher_KiemTraTruocKhiTaoDon", "Kiểm tra voucher theo user/cart/phí vận chuyển."),
    ("reviewApi.create", "POST /products/{id}/reviews", "ReviewsController.Create", "DANHGIASANPHAM, DONHANG, CHITIET_DONHANG", "Kiểm tra quyền đánh giá rồi lưu review Pending."),
    ("favoriteApi.add/remove", "POST/DELETE /favorites/{productId}", "FavoritesController", "YEUTHICH, SANPHAM", "Thêm/xóa yêu thích theo user hiện tại."),
    ("userApi address APIs", "GET/POST/PUT/DELETE /users/me/addresses", "UsersController", "NGUOIDUNG_DIACHI", "CRUD địa chỉ nhận hàng."),
    ("paymentApi.getByOrder", "GET /payments/order/{orderId}", "PaymentsController.GetByOrder -> PaymentRepository", "THANHTOAN, DONHANG", "Lấy danh sách giao dịch thanh toán của đơn."),
    ("contentApi.createContactRequest", "POST /content/contact-requests", "ContentController.CreateContactRequest", "LIENHE_YEUCAU, SANPHAM", "Lưu yêu cầu liên hệ/tư vấn."),
]


DB_TABLES = [
    ("NGUOIDUNG", "AuthService, OrderService, PaymentService, CatalogService review", "Tài khoản, profile, user id trong đơn/review/payment."),
    ("NGUOIDUNG_DIACHI", "AuthService", "Địa chỉ giao/nhận hàng user lưu trong Account/Checkout."),
    ("SANPHAM", "CatalogService, OrderService", "Nguồn catalog; list/detail đọc ở CatalogService, cart/order đọc lại để validate."),
    ("BIENSANPHAM", "CatalogService, OrderService", "SKU/biến thể, tồn kho, giá ghi đè hiện tại."),
    ("ANHSANPHAM", "CatalogService, OrderService", "Ảnh sản phẩm/biến thể cho card/detail/cart."),
    ("DANHMUC, HANGXE, DONGXE", "CatalogService", "Bộ lọc và tên hiển thị."),
    ("GIOHANG, CHITIET_GIOHANG", "OrderService", "Giỏ hàng active và dòng giỏ."),
    ("DONHANG, CHITIET_DONHANG", "OrderService, CatalogService review", "Đơn hàng và dòng đơn; CatalogService đọc để kiểm tra quyền review."),
    ("DONHANG_LICHSU_TRANGTHAI", "OrderService", "Lịch sử trạng thái đơn."),
    ("TONKHO_GIUCHO", "OrderService, PaymentService", "Giữ tồn cho đơn chưa hoàn tất hoặc đang chờ thanh toán."),
    ("THANHTOAN", "OrderService, PaymentService", "Giao dịch thanh toán hoặc yêu cầu thanh toán."),
    ("HOSO_TRAGOP", "OrderService", "Thông tin hồ sơ trả góp khi checkout."),
    ("YEUCAU_HOANTIEN", "OrderService", "Yêu cầu hoàn tiền từ OrderDetailPage."),
    ("VOUCHER, VOUCHER_NGUOIDUNG, DONHANG_VOUCHER", "OrderService", "Định nghĩa voucher, voucher đã lưu, voucher đã áp vào đơn."),
    ("VOUCHER_SANPHAM, VOUCHER_DANHMUC, VOUCHER_HANGXE", "OrderService raw SQL", "Scope voucher theo sản phẩm/danh mục/hãng."),
    ("DANHGIASANPHAM", "CatalogService", "Đánh giá sản phẩm."),
    ("YEUTHICH", "CatalogService", "Sản phẩm yêu thích."),
    ("FAQ, LIENHE_YEUCAU", "CatalogService", "FAQ public và yêu cầu liên hệ."),
    ("MATKHAU_DATLAI", "AuthService", "Token reset password."),
]


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Luồng chức năng Frontend -> Backend -> Database")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Bản tiếng Việt chi tiết: giải thích code, chức năng của code, API contract, backend và database. Phạm vi: Frontend khách hàng, không gồm FrontendAdmin.")

    add_note(
        doc,
        "Mục tiêu",
        "Tài liệu này không chỉ liệt kê màn hình. Mỗi phần giải thích code nhận dữ liệu gì, biến đổi gì, gọi file/hàm nào, gửi endpoint nào, backend xử lý ở đâu và database nào bị đọc/ghi.",
    )

    add_heading(doc, "1. Kiến trúc tổng quan", 1)
    add_bullets(doc, [
        "React page/component không nên gọi axios trực tiếp; luồng chuẩn là page -> context hoặc services/api.js -> httpClient.js -> gateway -> backend.",
        "api.js là lớp hợp đồng giữa UI và backend: đổi tên field, đổi query, chọn endpoint, rồi normalize response.",
        "normalizers.js và productMappers.js giúp UI không phụ thuộc backend trả camelCase, PascalCase hay tên tiếng Việt.",
        "Backend tính dữ liệu quan trọng như giá cuối, tổng tiền, voucher, tồn kho và trạng thái đơn. Frontend chỉ preview và hiển thị.",
    ])

    add_code_block(doc, "Pipeline request", [
        "Page/Component",
        "  -> Context hoặc services/api.js",
        "  -> normalizers.js/productMappers.js",
        "  -> httpClient.js gắn Authorization bearer",
        "  -> Vite proxy /api -> API Gateway localhost:5000",
        "  -> Controller",
        "  -> Service/Repository",
        "  -> DbContext/raw SQL/stored procedure",
    ])

    add_heading(doc, "2. Giải thích code Frontend theo lớp", 1)
    add_table(doc, ["File", "Chức năng code", "Giải thích chi tiết"], FRONTEND_CODE, [2300, 1900, 5160], font_size=8.0)

    add_heading(doc, "3. Giải thích code theo màn hình", 1)
    add_table(doc, ["Màn hình/file", "Code làm gì", "Tác động tới backend/database"], PAGE_CODE, [2100, 3900, 3360], font_size=7.7)

    add_heading(doc, "4. Contract API trong api.js", 1)
    add_table(doc, ["API helper", "Input từ UI", "Request gửi backend", "Kết quả/side effect"], API_CONTRACTS, [1600, 2300, 3200, 2260], font_size=7.5)

    add_heading(doc, "5. Route khách hàng và service liên quan", 1)
    add_table(doc, ["Route", "Frontend entry", "Chức năng màn hình", "API helper", "Service"], ROUTES, [1000, 1700, 2600, 2500, 1560], font_size=7.5)

    add_heading(doc, "6. Giải thích code Backend", 1)
    add_table(doc, ["File backend", "Chức năng code", "Giải thích chi tiết"], BACKEND_CODE, [2500, 1900, 4960], font_size=7.8)

    add_heading(doc, "7. Endpoint -> backend -> database", 1)
    add_table(doc, ["Frontend API", "Endpoint", "Backend xử lý", "Bảng/SP chính", "Code thực hiện gì"], ENDPOINTS, [1450, 1500, 2300, 2500, 1610], font_size=7.3)

    add_heading(doc, "8. Pseudo-code các luồng quan trọng", 1)
    add_code_block(doc, "Checkout tạo đơn", [
        "CheckoutPage submit form",
        "  đọc cart từ CartContext",
        "  gọi voucherApi.validate nếu có mã",
        "  gọi orderApi.getShippingQuote để preview phí",
        "  gọi orderApi.create(payload)",
        "OrderService.CreateFromCart",
        "  load GIOHANG + CHITIET_GIOHANG",
        "  validate SANPHAM/BIENSANPHAM/tồn kho",
        "  tính lại subtotal, discount, shipping",
        "  insert DONHANG + CHITIET_DONHANG",
        "  ghi DONHANG_LICHSU_TRANGTHAI",
        "  tạo TONKHO_GIUCHO, THANHTOAN, HOSO_TRAGOP nếu có",
        "  ghi nhận voucher bằng stored procedure",
    ])
    add_code_block(doc, "Thêm vào giỏ hàng", [
        "ProductCard/ProductDetailPage",
        "  build payload { productId, variantId, quantity }",
        "cartApi.addItem",
        "  POST /cart/items { maSanPham, maBienSanPham, soLuong }",
        "OrderService",
        "  get/create GIOHANG active",
        "  validate sản phẩm, biến thể, tồn kho",
        "  insert/update CHITIET_GIOHANG",
        "  trả CartDto mới để CartContext cập nhật count/subtotal",
    ])

    add_heading(doc, "9. Bảng database và vai trò", 1)
    add_table(doc, ["Bảng", "Service đọc/ghi", "Vai trò trong luồng Frontend"], DB_TABLES, [2600, 2300, 4460], font_size=7.8)

    add_heading(doc, "10. Stored procedure/raw SQL quan trọng", 1)
    add_table(
        doc,
        ["Stored procedure/raw SQL", "Luồng gọi", "Chức năng"],
        [
            ("sp_Voucher_KiemTraTruocKhiTaoDon", "voucherApi.validate và orderApi.create", "Kiểm tra voucher theo user, giỏ hàng, mã voucher, phí vận chuyển; trả hợp lệ/lý do/số tiền giảm."),
            ("sp_Voucher_GhiNhanSuDung", "Tạo đơn thành công", "Ghi nhận voucher đã dùng vào đơn và tăng usage."),
            ("sp_Voucher_HuySuDungTheoDon", "Hủy đơn", "Rollback voucher đã ghi nhận khi đơn bị hủy."),
            ("sp_TONKHO_ApDungBienDong", "Tạo/hủy/confirm đơn và nghiệp vụ tồn", "Cập nhật tồn/giữ chỗ theo MaSanPham/MaBienSanPham và ghi log tồn kho."),
            ("Raw SQL VOUCHER_SANPHAM/DANHMUC/HANGXE", "voucherApi.getApplicable", "Lọc voucher theo scope sản phẩm, danh mục hoặc hãng xe."),
        ],
        [2800, 2300, 4260],
        font_size=8.0,
    )

    add_heading(doc, "11. Ghi chú khi sửa code", 1)
    add_bullets(doc, [
        "Nếu UI hiển thị sai, kiểm tra theo thứ tự: component state -> api.js -> normalizers/productMappers -> response backend -> query database.",
        "Không nên để page tự gọi axios trực tiếp nếu đã có helper trong api.js, vì sẽ phá contract mapping tập trung.",
        "Không dùng giá/tổng tiền frontend làm dữ liệu cuối cùng cho đơn hàng. Backend OrderService phải tính lại và snapshot vào DONHANG/CHITIET_DONHANG.",
        "Khi backend đổi field, nên bổ sung alias trong mapper thay vì sửa rải rác ở từng component.",
        "FrontendAdmin không nằm trong phạm vi tài liệu này, dù một số controller backend có endpoint admin trong cùng file.",
    ])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("ShowRoomDB - tài liệu luồng Frontend khách hàng").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
